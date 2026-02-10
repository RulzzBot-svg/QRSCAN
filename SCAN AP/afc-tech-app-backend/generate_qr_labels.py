#!/usr/bin/env python3
"""Generate QR code PNGs (and optional Word doc) for AHUs.

Usage examples:
  python generate_qr_labels.py --hospital "Foothill" --out qr_codes --doc
  python generate_qr_labels.py --all --base-url "https://example.com/FilterInfo/" --out out/qrs

This is a restored utility similar to the previous `qrutils.py`.
"""
import os
import argparse
import qrcode
from app import create_app
from db import db
from models import AHU, Hospital

try:
    from docx import Document
    from docx.shared import Inches
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False


def ensure_dir(path):
    os.makedirs(path, exist_ok=True)


def make_qr_image(url, out_path, box_size=6):
    qr = qrcode.QRCode(
        version=2,
        error_correction=qrcode.constants.ERROR_CORRECT_Q,
        box_size=box_size,
        border=4,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save(out_path)


def gen_for_hospital(app, hospital, out_dir, base_url, make_doc=False):
    ensure_dir(out_dir)
    ahus = AHU.query.filter_by(hospital_id=hospital.id).order_by(AHU.id).all()

    doc = None
    if make_doc:
        if not HAVE_DOCX:
            print("python-docx not available; skipping Word document creation")
            make_doc = False
        else:
            doc = Document()
            doc.add_heading(f'QR Codes for {hospital.name}', 0)

    for ahu in ahus:
        ahu_id = ahu.id
        label = ahu.name or str(ahu_id)
        url = f"{base_url.rstrip('/')}/{ahu_id}"
        filename = f"{ahu_id}.png"
        out_path = os.path.join(out_dir, filename)
        make_qr_image(url, out_path)
        print(f"Wrote {out_path} → {url}")

        if make_doc and doc is not None:
            doc.add_heading(f'AHU: {label}', level=1)
            if ahu.location:
                doc.add_paragraph(f'Location: {ahu.location}')
            try:
                doc.add_picture(out_path, width=Inches(2))
            except Exception as e:
                print(f"Failed to add picture to doc: {e}")
            doc.add_paragraph('')

    if make_doc and doc is not None:
        doc_path = os.path.join(out_dir, f"{hospital.name.replace(' ', '_')}_QRs.docx")
        doc.save(doc_path)
        print(f"Saved Word doc: {doc_path}")


def main():
    parser = argparse.ArgumentParser(description="Generate QR codes for AHUs")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--hospital", "-H", help="Hospital name to generate for (exact match)")
    group.add_argument("--all", action="store_true", help="Generate for all hospitals")
    parser.add_argument("--out", "-o", default="qr_codes", help="Output directory")
    parser.add_argument("--base-url", "-b", default="https://qrscan-lyart.vercel.app/FilterInfo/", help="Base URL to embed in QR (append AHU id)")
    parser.add_argument("--doc", action="store_true", help="Also produce a Word document with embedded images (requires python-docx)")

    args = parser.parse_args()

    app = create_app()
    with app.app_context():
        ensure_dir(args.out)
        if args.all:
            hospitals = Hospital.query.order_by(Hospital.name).all()
            for h in hospitals:
                out_dir = os.path.join(args.out, h.name.replace(' ', '_'))
                print(f"Generating for hospital: {h.name} → {out_dir}")
                gen_for_hospital(app, h, out_dir, args.base_url, make_doc=args.doc)
        else:
            h = Hospital.query.filter_by(name=args.hospital).first()
            if not h:
                print(f"Hospital '{args.hospital}' not found. Available:")
                for hh in Hospital.query.order_by(Hospital.name).all():
                    print(f"  - {hh.name}")
                return
            gen_for_hospital(app, h, args.out, args.base_url, make_doc=args.doc)


if __name__ == '__main__':
    main()
