import qrcode
import os
import sys
from app import create_app
from db import db
from models import AHU, Hospital
from docx import Document
from docx.shared import Inches

app = create_app()

with app.app_context():
    if len(sys.argv) < 2:
        print("Usage: python qrutils.py <hospital_name>")
        print("Available hospitals:")
        hospitals = Hospital.query.all()
        for h in hospitals:
            print(f"  - {h.name}")
        sys.exit(1)

    hospital_name = sys.argv[1]
    hospital = Hospital.query.filter_by(name=hospital_name).first()
    if not hospital:
        print(f"Hospital '{hospital_name}' not found.")
        sys.exit(1)

    ahus = AHU.query.filter_by(hospital_id=hospital.id).all()

    # Create Word document
    doc = Document()
    doc.add_heading(f'QR Codes for {hospital_name}', 0)

    for ahu in ahus:
        ahu_id = ahu.id
        url = f"https://qrscan-lyart.vercel.app/FilterInfo/{ahu_id}"

        qr = qrcode.QRCode(
            version=2,
            error_correction=qrcode.constants.ERROR_CORRECT_Q,
            box_size=10,
            border=4
        )

        qr.add_data(url)
        qr.make(fit=True)

        img = qr.make_image(fill_color="black", back_color="white")
        img_path = f"qr_codes/{ahu_id}.png"
        img.save(img_path)
        print(f"QR generated for {ahu_id}")

        # Add to Word doc
        doc.add_heading(f'AHU: {ahu.name or ahu_id}', level=1)
        if ahu.location:
            doc.add_paragraph(f'Location: {ahu.location}')
        doc.add_picture(img_path, width=Inches(2))
        doc.add_paragraph('')  # Spacer

    doc_path = f"qr_codes/{hospital_name.replace(' ', '_')}_QRs.docx"
    doc.save(doc_path)
    print(f"Word document saved: {doc_path}")

print(f"All QR codes and Word document generated for {hospital_name}!")